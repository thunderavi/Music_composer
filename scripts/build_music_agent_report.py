from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("Music_Composition_Agent_One_Page_Report.docx")


COLORS = {
    "ink": "17202A",
    "blue": "1F4E79",
    "green": "1E6B4E",
    "purple": "5B3F8C",
    "orange": "A65F00",
    "gray": "EEF2F5",
    "line": "D6DEE6",
    "dark": "111827",
    "soft": "F7FAFC",
    "caution": "FFF4D6",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D6DEE6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
        borders.append(tag)
    tc_pr.append(borders)


def set_cell_margins(cell, top=65, bottom=65, left=90, right=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def style_run(run, size=8.7, bold=False, color="17202A"):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def para(cell, text="", size=8.7, bold=False, color="17202A", after=0, align=None):
    p = cell.paragraphs[0] if len(cell.paragraphs) == 1 and not cell.paragraphs[0].text else cell.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if align:
        p.alignment = align
    r = p.add_run(text)
    style_run(r, size=size, bold=bold, color=color)
    return p


def add_label_value(cell, label, value, color="1F4E79"):
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label.upper() + "\n")
    style_run(r, size=6.3, bold=True, color=color)
    r = p.add_run(value)
    style_run(r, size=8.3, bold=True, color="17202A")


def add_chip(cell, text, fill, color="17202A"):
    p = cell.paragraphs[0] if not cell.paragraphs[0].text else cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    style_run(r, size=7.2, bold=True, color=color)
    set_cell_shading(cell, fill)


def add_heading(doc, text, color="1F4E79"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    style_run(r, size=10.2, bold=True, color=color)
    return p


def add_bullets(cell, items, size=7.7):
    for item in items:
        p = cell.add_paragraph(style=None)
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        p.paragraph_format.space_after = Pt(0.5)
        r = p.add_run("- ")
        style_run(r, size=size, bold=True)
        r = p.add_run(item)
        style_run(r, size=size)


def format_table(table, fill_header=None):
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
            set_cell_margins(cell)
            if row_idx == 0 and fill_header:
                set_cell_shading(cell, fill_header)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.45)
section.bottom_margin = Inches(0.45)
section.left_margin = Inches(0.55)
section.right_margin = Inches(0.55)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(8.4)
normal.paragraph_format.space_after = Pt(2)
normal.paragraph_format.line_spacing = 1.02

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(1)
r = title.add_run("Music Composition Agent")
style_run(r, size=17.2, bold=True, color="111827")
r = title.add_run("  |  NVIDIA NIM Solution + Market Report")
style_run(r, size=10.2, bold=True, color="1F4E79")

sub = doc.add_paragraph()
sub.paragraph_format.space_after = Pt(4)
r = sub.add_run(
    "One-page brief covering problem fit, solution architecture, technical approach, market comparison, and execution strategy."
)
style_run(r, size=7.8, color="4B5563")

meta = doc.add_table(rows=1, cols=4)
set_table_width(meta, [1.55, 1.7, 1.8, 2.1])
format_table(meta)
for cell, label, value in zip(
    meta.rows[0].cells,
    ["Problem", "AI Provider", "Backend/API", "Output"],
    [
        "Fast editable music drafts",
        "NVIDIA NIM only",
        "FastAPI /api/v1 on :5050",
        "Chords, melody, lyrics, MIDI",
    ],
):
    add_label_value(cell, label, value)

doc.add_paragraph().paragraph_format.space_after = Pt(0)

lead = doc.add_table(rows=1, cols=1)
set_table_width(lead, [7.4])
format_table(lead)
cell = lead.rows[0].cells[0]
set_cell_shading(cell, "F3F8FF")
p = cell.paragraphs[0]
p.paragraph_format.space_after = Pt(0)
r = p.add_run("Positioning: ")
style_run(r, size=8.6, bold=True, color="1F4E79")
r = p.add_run(
    "Unlike black-box audio generators, this project creates editable musical structure: chord progressions, symbolic melody, lyrics, draft storage, and MIDI/notation export through a transparent NVIDIA NIM + FastAPI + React pipeline."
)
style_run(r, size=8.3, color="17202A")

add_heading(doc, "Solution Architecture & Working Flow")
flow = doc.add_table(rows=2, cols=6)
set_table_width(flow, [1.08, 1.16, 1.2, 1.18, 1.28, 1.5])
format_table(flow)
flow_labels = [
    ("User Input", "Style, mood, theme, key, BPM, bars"),
    ("React UI", "Editable form and draft workspace"),
    ("FastAPI", "Routes: provider, compose, drafts, exports"),
    ("NVIDIA NIM", "Chat completion returns JSON only"),
    ("Validation", "Pydantic + music checks"),
    ("Storage/Export", "SQLite, music21 MIDI, notation txt"),
]
fills = ["DFF7F4", "E5F0FF", "E8F6EC", "F1E8FF", "FFF0DA", "F2F4F7"]
for idx, (label, detail) in enumerate(flow_labels):
    add_chip(flow.rows[0].cells[idx], label, fills[idx])
    p = flow.rows[1].cells[idx].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(detail)
    style_run(r, size=6.7, color="374151")

add_heading(doc, "Technical Approach")
tech = doc.add_table(rows=1, cols=3)
set_table_width(tech, [2.45, 2.45, 2.5])
format_table(tech)
heads = ["Generation", "Validation + Control", "User Execution"]
items = [
    ["Prompt builder creates strict JSON schema", "NIM base URL normalized to /v1", "No OpenAI/Gemini fallback"],
    ["Pydantic validates response shape", "Music checks chord symbols and note pitches", "Errors/warnings shown before export"],
    ["Edit chords, melody, lyrics in UI", "Save draft through SQLite", "Export MIDI via music21"],
]
for i, cell in enumerate(tech.rows[0].cells):
    p = cell.paragraphs[0]
    r = p.add_run(heads[i])
    style_run(r, size=8.2, bold=True, color="1F4E79")
    add_bullets(cell, items[i], size=7.25)

add_heading(doc, "Market Landscape")
market = doc.add_table(rows=1, cols=5)
set_table_width(market, [1.0, 1.55, 1.55, 1.55, 1.35])
format_table(market, fill_header="E8EEF5")
headers = ["Tool", "How it works", "Strength", "Gap vs our need", "Source"]
for idx, h in enumerate(headers):
    p = market.rows[0].cells[idx].paragraphs[0]
    r = p.add_run(h)
    style_run(r, size=7.2, bold=True, color="17202A")
rows = [
    ["Suno", "Prompt to full audio song", "Vocals + polished demos", "Less symbolic editing/control", "suno.com"],
    ["Udio", "Prompt/lyrics to song audio", "Realistic creative audio", "Audio-first, less transparent", "udio.com"],
    ["Mubert", "API generates royalty-free tracks", "Good developer/API use", "Background music focus", "mubert.com"],
    ["Boomy", "Choose style, generate/publish", "Beginner publishing flow", "Closed platform internals", "boomy.com"],
    ["Our Agent", "NIM -> JSON -> validate -> edit/export", "Transparent, editable, testable", "MVP audio quality, no vocals yet", "Local project"],
]
for row in rows:
    cells = market.add_row().cells
    for idx, text in enumerate(row):
        p = cells[idx].paragraphs[0]
        r = p.add_run(text)
        style_run(r, size=6.65, bold=(idx == 0), color="17202A")

add_heading(doc, "Why Our Project Wins for Evaluation")
win = doc.add_table(rows=1, cols=2)
set_table_width(win, [3.7, 3.7])
format_table(win)
for idx, (heading, bullets) in enumerate(
    [
        (
            "Better for hackathon/portfolio",
            [
                "Shows full-stack execution, not just prompt output",
                "Demonstrates agentic pipeline: prompt, validate, store, export",
                "Easy to explain architecture and evaluation criteria",
            ],
        ),
        (
            "Better for musicians learning/composing",
            [
                "Outputs editable chords, melody notes, and lyrics",
                "MIDI can be imported into DAWs for real production",
                "Clear disclaimer and rights-review note included",
            ],
        ),
    ]
):
    cell = win.rows[0].cells[idx]
    set_cell_shading(cell, "F8FBF8" if idx == 0 else "FFF9EE")
    p = cell.paragraphs[0]
    r = p.add_run(heading)
    style_run(r, size=8.2, bold=True, color="1E6B4E" if idx == 0 else "A65F00")
    add_bullets(cell, bullets, size=7.25)

foot = doc.add_paragraph()
foot.paragraph_format.space_before = Pt(4)
foot.paragraph_format.space_after = Pt(0)
r = foot.add_run(
    "Execution depth: React + Vite frontend, FastAPI backend, NVIDIA NIM-only generation, .env security, SQLite persistence, Pydantic validation, music21 MIDI export, tests, and Docker-ready setup. Disclaimer: generated music may resemble existing works; review before commercial use."
)
style_run(r, size=6.9, color="4B5563")

doc.save(OUT)
print(OUT.resolve())
